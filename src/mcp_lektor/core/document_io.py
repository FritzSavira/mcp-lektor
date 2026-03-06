"""Read .docx files and convert them to a structured DocumentStructure."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument  # python-docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from mcp_lektor.core.models import (
    DocumentParagraph,
    DocumentStructure,
    ParagraphType,
    RunFormatting,
    TextColor,
    TextRun,
)
from mcp_lektor.core.run_normalizer import normalize_runs

logger = logging.getLogger(__name__)

# ------------------------------------------------------------------ #
# Public API                                                          #
# ------------------------------------------------------------------ #


def parse_docx(file_path: str | Path) -> DocumentStructure:
    """Parse a .docx file and return a fully populated DocumentStructure."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {path.suffix}")

    doc = DocxDocument(str(path))
    paragraphs: list[DocumentParagraph] = []

    for idx, para in enumerate(doc.paragraphs):
        raw_runs = [_build_text_run(run) for run in para.runs]
        runs = normalize_runs(raw_runs)

        has_content_runs = [r for r in runs if r.text.strip()]
        is_placeholder = bool(has_content_runs) and all(
            r.is_placeholder for r in has_content_runs
        )

        doc_para = DocumentParagraph(
            index=idx,
            paragraph_type=_classify_paragraph(para),
            style_name=para.style.name if para.style else None,
            heading_level=_get_heading_level(para),
            runs=runs,
            is_placeholder_paragraph=is_placeholder,
        )
        paragraphs.append(doc_para)

    total_words = sum(len(p.plain_text.split()) for p in paragraphs)
    placeholder_paras = [p for p in paragraphs if p.is_placeholder_paragraph]

    return DocumentStructure(
        filename=path.name,
        paragraphs=paragraphs,
        total_paragraphs=len(paragraphs),
        total_words=total_words,
        placeholder_count=len(placeholder_paras),
        placeholder_locations=[f"Paragraph {p.index}" for p in placeholder_paras],
    )


# ------------------------------------------------------------------ #
# Run helpers                                                         #
# ------------------------------------------------------------------ #


def _build_text_run(run: Run) -> TextRun:
    formatting = _extract_run_formatting(run)
    return TextRun(
        text=run.text,
        formatting=formatting,
        is_placeholder=_is_placeholder(run, formatting),
    )


def _extract_run_formatting(run: Run) -> RunFormatting:
    color = _extract_color(run)
    return RunFormatting(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        strike=bool(run.font.strike),
        font_name=run.font.name,
        font_size=run.font.size.pt if run.font.size else None,
        color=color,
        highlight=str(run.font.highlight_color) if run.font.highlight_color else None,
        style_name=run.style.name if run.style else None,
    )


def _extract_color(run: Run) -> Optional[TextColor]:
    """Try to read an explicit RGB color from the run."""
    # python-docx exposes color.rgb as an RGBColor when set explicitly
    rgb = run.font.color.rgb
    if rgb is not None:
        return TextColor(r=rgb[0], g=rgb[1], b=rgb[2])

    # Fallback: inspect the XML for <w:color w:val="..."/>
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        color_elem = rpr.find(qn("w:color"))
        if color_elem is not None:
            val = color_elem.get(qn("w:val"))
            if val and len(val) == 6:
                try:
                    r = int(val[0:2], 16)
                    g = int(val[2:4], 16)
                    b = int(val[4:6], 16)
                    return TextColor(r=r, g=g, b=b)
                except ValueError:
                    pass
    return None


def _is_placeholder(run: Run, formatting: RunFormatting) -> bool:
    """A run counts as placeholder when it has red text."""
    if formatting.color and formatting.color.is_red:
        return True
    return False


# ------------------------------------------------------------------ #
# Paragraph helpers                                                   #
# ------------------------------------------------------------------ #

_HEADING_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
_LIST_STYLES = {"List Paragraph", "List Bullet", "List Number"}


def _classify_paragraph(para: Paragraph) -> ParagraphType:
    style_name = para.style.name if para.style else ""
    if _HEADING_RE.match(style_name):
        return ParagraphType.HEADING
    if style_name in _LIST_STYLES:
        return ParagraphType.LIST_ITEM
    return ParagraphType.BODY


def _get_heading_level(para: Paragraph) -> Optional[int]:
    style_name = para.style.name if para.style else ""
    m = _HEADING_RE.match(style_name)
    return int(m.group(1)) if m else None


# ------------------------------------------------------------------ #
# Write-back (add to existing document_io.py)                        #
# ------------------------------------------------------------------ #


def write_corrected_document(
    input_path: str | Path,
    output_path: str | Path,
    corrections: list[dict],
    author: str = "MCP-Lektor",
    decisions: Optional[dict[int, str]] = None,
) -> Path:
    """Apply corrections as Track Changes and save to output_path.

    Args:
        input_path: Path to the original .docx file.
        output_path: Path where the corrected .docx will be saved.
        corrections: List of correction dicts from the proofreading engine.
        author: Revision author name shown in Track Changes.
        decisions: Optional mapping of correction index → "accept"/"reject"/"edit".

    Returns:
        The output path as a Path object.
    """
    from mcp_lektor.core.openxml_writer import (
        _save_comments_part,
        apply_corrections_to_document,
    )

    input_path = Path(input_path)
    output_path = Path(output_path)

    doc = DocxDocument(str(input_path))
    apply_corrections_to_document(doc, corrections, author, decisions)
    # Perform a safety check on the XML before saving if needed, 
    # but python-docx's .save() is usually what does the XML assembly.
    _save_comments_part(doc)
    doc.save(str(output_path))

    # Verify structural integrity of the newly saved file
    from mcp_lektor.utils.xml_validator import validate_docx_structure
    try:
        validate_docx_structure(output_path)
    except Exception as e:
        logger.error(f"Generated .docx is invalid: {e}")
        # We might want to raise an error here to prevent delivering corrupt files
        raise ValueError(f"Failed to generate a valid .docx: {e}")

    logger.info("Corrected document saved and validated at %s", output_path)
    return output_path
