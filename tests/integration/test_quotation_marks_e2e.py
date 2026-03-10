"""End-to-end test for consolidated quotation mark logic."""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument
from lxml import etree

from mcp_lektor.core.models import (
    DocumentParagraph,
    DocumentStructure,
    ParagraphType,
    RunFormatting,
    TextRun,
)
from mcp_lektor.core.openxml_writer import apply_corrections_to_document
from mcp_lektor.core.proofreading_engine import ProofreadingEngine


def _make_structure_from_doc(doc: DocxDocument) -> DocumentStructure:
    """Helper to convert a docx.Document to our DocumentStructure."""
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        runs = [TextRun(text=r.text, formatting=RunFormatting()) for r in p.runs]
        paragraphs.append(
            DocumentParagraph(
                index=i,
                paragraph_type=ParagraphType.BODY,
                runs=runs,
            )
        )
    return DocumentStructure(
        filename="test.docx",
        paragraphs=paragraphs,
        total_paragraphs=len(paragraphs),
    )

@pytest.mark.asyncio
async def test_quotation_marks_to_word_export():
    # 1. Create a document with straight quotes
    doc = DocxDocument()
    # Paragraph with opening and closing straight quotes
    doc.add_paragraph('Er sagte "Hallo".')

    # 2. Run ProofreadingEngine
    structure = _make_structure_from_doc(doc)
    engine = ProofreadingEngine()
    result = await engine.proofread(structure)

    # Verify we got corrections for quotation marks
    assert result.total_corrections >= 2

    # 3. Apply corrections to document
    # Convert ProposedCorrection objects to dicts for apply_corrections_to_document
    corrections_data = []
    for c in result.corrections:
        corrections_data.append({
            "paragraph_index": c.paragraph_index,
            "char_offset_start": c.char_offset_start,
            "original_text": c.original_text,
            "suggested_text": c.suggested_text,
            "category": c.category.value,
            "explanation": c.explanation,
        })

    corrected_doc = apply_corrections_to_document(doc, corrections_data, author="Lektor")

    # 4. Verify XML output
    para_xml = etree.tostring(corrected_doc.paragraphs[0]._element, encoding="unicode")

    # Check for track changes
    assert "w:del" in para_xml
    assert "w:ins" in para_xml

    # Check for correct typographic marks
    assert "„" in para_xml # Opening
    assert "“" in para_xml # Closing

    # Check for comments
    assert "commentRangeStart" in para_xml
    assert "commentReference" in para_xml

    # Verify the explanation is in the comments (if we can find the comment element)
    # The comments are stored in corrected_doc._comments_element
    comments_xml = etree.tostring(corrected_doc._comments_element, encoding="unicode")
    assert "Deutsche öffnende Anführungszeichen sind „ (unten)." in comments_xml
    assert "Deutsche schließende Anführungszeichen sind “ (oben)." in comments_xml
