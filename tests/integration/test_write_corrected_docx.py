"""Integration tests for the write_corrected_docx tool."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from mcp_lektor.core.session_manager import session_manager
from mcp_lektor.core.models import (
    ProofreadingResult,
    ProposedCorrection,
    CorrectionCategory,
    ConfidenceLevel
)
from mcp_lektor.tools.write_corrected_docx import write_corrected_docx


def _create_test_docx(
    path: Path, text: str = "Dies ist ein Testtext mit Fehler."
) -> Path:
    """Create a simple test .docx file."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _setup_session(session_id: str, file_path: Path) -> str:
    """Set up a session with corrections. 
    Uses session_manager.create_session to ensure timestamps are set.
    """
    corrections = [
        ProposedCorrection(
            id="C-001",
            paragraph_index=0,
            run_index=0,
            char_offset_start=27,
            char_offset_end=33,
            original_text="Fehler",
            suggested_text="Erfolg",
            category=CorrectionCategory.SPELLING,
            confidence=ConfidenceLevel.HIGH,
            explanation="Testkorrektur",
        ),
    ]
    result = ProofreadingResult(
        document_filename=file_path.name,
        total_corrections=len(corrections),
        corrections=corrections
    )
    # Clear and recreate to avoid ID collision if needed, 
    # but here we just use the manager properly.
    data = {
        "file_path": str(file_path),
        "proofreading_result": result,
    }
    # Hack to use a specific session_id for tests that expect it
    sid = session_manager.create_session(data)
    # If the test really needs a specific ID, we'd have to swap it in the dict
    session_manager._sessions[session_id] = session_manager._sessions.pop(sid)
    return session_id


class TestWriteCorrectedDocx:
    """Integration tests for the write_corrected_docx MCP tool."""

    @pytest.mark.asyncio
    async def test_apply_all_mode(self, tmp_path: Path) -> None:
        docx_path = _create_test_docx(tmp_path / "test.docx")
        session_id = "test-session-1"
        _setup_session(session_id, docx_path)

        result_raw = await write_corrected_docx(
            session_id=session_id,
            apply_all=True,
        )
        result = json.loads(result_raw)

        assert result.get("status") == "success", f"Error: {result.get('error')}"
        assert result["corrections_applied"] == 1
        assert Path(result["output_path"]).exists()

    @pytest.mark.asyncio
    async def test_selective_decisions(self, tmp_path: Path) -> None:
        docx_path = _create_test_docx(tmp_path / "test2.docx")
        session_id = "test-session-2"
        _setup_session(session_id, docx_path)

        # Use index 0 for the first correction
        decisions = json.dumps({"0": "reject"})
        result_raw = await write_corrected_docx(
            session_id=session_id,
            decisions=decisions,
            apply_all=False,
        )
        result = json.loads(result_raw)

        assert result.get("status") == "success"
        assert result["corrections_applied"] == 0

    @pytest.mark.asyncio
    async def test_missing_session(self) -> None:
        result_raw = await write_corrected_docx(session_id="nonexistent")
        result = json.loads(result_raw)
        assert "error" in result
        assert "not found" in result["error"]

    @pytest.mark.asyncio
    async def test_no_corrections(self, tmp_path: Path) -> None:
        docx_path = _create_test_docx(tmp_path / "test3.docx")
        session_id = "empty-session"
        result_obj = ProofreadingResult(
            document_filename=docx_path.name,
            total_corrections=0,
            corrections=[]
        )
        session_manager.create_session({
            "file_path": str(docx_path),
            "proofreading_result": result_obj,
        })
        # Need to fix the ID for the test
        last_id = list(session_manager._sessions.keys())[-1]
        session_manager._sessions[session_id] = session_manager._sessions.pop(last_id)

        result_raw = await write_corrected_docx(session_id=session_id)
        result = json.loads(result_raw)
        assert result.get("status") == "success"
        assert "nothing to write" in result["message"]

    @pytest.mark.asyncio
    async def test_output_opens_in_docx(self, tmp_path: Path) -> None:
        docx_path = _create_test_docx(tmp_path / "test4.docx")
        session_id = "test-session-4"
        _setup_session(session_id, docx_path)

        result_raw = await write_corrected_docx(session_id=session_id)
        result = json.loads(result_raw)
        assert result.get("status") == "success"

        # Verify the output is a valid .docx
        output = DocxDocument(result["output_path"])
        assert len(output.paragraphs) > 0
