"""Unit tests for the OpenXML writer module."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from lxml import etree

from mcp_lektor.core.openxml_writer import (
    WORD_NS,
    add_comment,
    apply_corrections_to_document,
    apply_track_change,
)


def _create_simple_docx(
    text: str = "Dies ist ein Testtext.", bold: bool = False
) -> DocxDocument:
    """Create a minimal .docx with one paragraph and one run."""
    doc = DocxDocument()
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    return doc


def _get_paragraph_element(doc: DocxDocument, index: int = 0) -> etree._Element:
    """Get the lxml element for a paragraph by index."""
    paragraphs = doc.element.body.findall(f".//{{{WORD_NS}}}p")
    return paragraphs[index]


class TestApplyTrackChange:
    """Tests for apply_track_change()."""

    def test_creates_del_and_ins_elements(self) -> None:
        doc = _create_simple_docx("Dies ist ein Fehler im Text.")
        para = _get_paragraph_element(doc)

        apply_track_change(
            paragraph_element=para,
            run_index=0,
            char_start=13,
            char_end=19,
            original_text="Fehler",
            replacement_text="Erfolg",
            author="Test",
            timestamp="2026-03-05T12:00:00Z",
            revision_id=1,
        )

        xml_str = etree.tostring(para, encoding="unicode")
        assert f"{{{WORD_NS}}}del" in xml_str or "w:del" in xml_str
        assert f"{{{WORD_NS}}}ins" in xml_str or "w:ins" in xml_str

    def test_preserves_surrounding_text(self) -> None:
        doc = _create_simple_docx("Anfang Fehler Ende")
        para = _get_paragraph_element(doc)

        apply_track_change(
            paragraph_element=para,
            run_index=0,
            char_start=7,
            char_end=13,
            original_text="Fehler",
            replacement_text="Erfolg",
            author="Test",
            timestamp="2026-03-05T12:00:00Z",
            revision_id=1,
        )

        xml_str = etree.tostring(para, encoding="unicode")
        # Before-text and after-text runs should exist
        assert "Anfang " in xml_str
        assert " Ende" in xml_str

    def test_formatting_preserved_after_track_change(self) -> None:
        doc = _create_simple_docx("Fetter Text hier.", bold=True)
        para = _get_paragraph_element(doc)

        apply_track_change(
            paragraph_element=para,
            run_index=0,
            char_start=7,
            char_end=11,
            original_text="Text",
            replacement_text="Satz",
            author="Test",
            timestamp="2026-03-05T12:00:00Z",
            revision_id=1,
        )

        # Check that rPr (bold) is preserved in the inserted run
        ins_elements = para.findall(f".//{{{WORD_NS}}}ins")
        assert len(ins_elements) > 0
        ins_run = ins_elements[0].find(f".//{{{WORD_NS}}}r")
        rpr = ins_run.find(f"{{{WORD_NS}}}rPr")
        assert rpr is not None

    def test_out_of_range_run_index_is_safe(self) -> None:
        doc = _create_simple_docx("Test.")
        para = _get_paragraph_element(doc)

        # Should not raise
        apply_track_change(
            paragraph_element=para,
            run_index=99,
            char_start=0,
            char_end=4,
            original_text="Test",
            replacement_text="Prüf",
            author="Test",
            timestamp="2026-03-05T12:00:00Z",
            revision_id=1,
        )


class TestAddComment:
    """Tests for add_comment()."""

    def test_creates_comment_range_markers(self) -> None:
        doc = _create_simple_docx("Kommentierter Text.")
        para = _get_paragraph_element(doc)

        add_comment(
            document=doc,
            paragraph_element=para,
            run_index=0,
            comment_text="Testkommentar",
            author="Test",
            timestamp="2026-03-05T12:00:00Z",
            comment_id=0,
        )

        xml_str = etree.tostring(para, encoding="unicode")
        assert "commentRangeStart" in xml_str
        assert "commentRangeEnd" in xml_str
        assert "commentReference" in xml_str


class TestApplyCorrectionsToDocument:
    """Tests for apply_corrections_to_document()."""

    def test_applies_all_corrections(self) -> None:
        doc = _create_simple_docx("Dies ist ein Fehler und ein Irrtum.")
        corrections = [
            {
                "paragraph_index": 0,
                "run_index": 0,
                "char_start": 13,
                "char_end": 19,
                "original_text": "Fehler",
                "replacement_text": "Erfolg",
                "category": "Rechtschreibung",
                "explanation": "Testkorrektur 1",
            },
        ]

        apply_corrections_to_document(doc, corrections, author="Test")

        xml_str = etree.tostring(doc.element.body, encoding="unicode")
        assert "Fehler" in xml_str  # in delText
        assert "Erfolg" in xml_str  # in ins run

    def test_respects_reject_decision(self) -> None:
        doc = _create_simple_docx("Ein Fehler hier.")
        corrections = [
            {
                "paragraph_index": 0,
                "run_index": 0,
                "char_start": 4,
                "char_end": 10,
                "original_text": "Fehler",
                "replacement_text": "Erfolg",
                "category": "Rechtschreibung",
                "explanation": "Abgelehnt",
            },
        ]

        apply_corrections_to_document(
            doc, corrections, author="Test", decisions={0: "reject"}
        )

        xml_str = etree.tostring(doc.element.body, encoding="unicode")
        # Rejected correction should NOT be applied
        assert "del" not in xml_str or "delText" not in xml_str

    def test_roundtrip_saves_and_opens(self, tmp_path: Path) -> None:
        doc = _create_simple_docx("Testdokument zum Speichern.")
        corrections = [
            {
                "paragraph_index": 0,
                "run_index": 0,
                "char_start": 0,
                "char_end": 13,
                "original_text": "Testdokument",
                "replacement_text": "Prüfdokument",
                "category": "Test",
                "explanation": "Roundtrip-Test",
            },
        ]

        apply_corrections_to_document(doc, corrections, author="Test")
        out = tmp_path / "roundtrip_test.docx"
        doc.save(str(out))

        # Should open without error
        reopened = DocxDocument(str(out))
        assert len(reopened.paragraphs) > 0
