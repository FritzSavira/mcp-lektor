"""Shared pytest fixtures for MCP Lektor tests."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from mcp_lektor.core.models import (
    DocumentParagraph,
    DocumentStructure,
    ParagraphType,
    RunFormatting,
    TextColor,
    TextRun,
)


@pytest.fixture
def sample_red_color() -> TextColor:
    """A red color that should be detected as placeholder."""
    return TextColor(r=255, g=0, b=0)


@pytest.fixture
def sample_black_color() -> TextColor:
    """A standard black color."""
    return TextColor(r=0, g=0, b=0)


@pytest.fixture
def sample_document_structure() -> DocumentStructure:
    """A minimal DocumentStructure for testing."""
    runs = [
        TextRun(
            text="Dies ist ein Testabsatz.",
            formatting=RunFormatting(),
        ),
    ]
    paragraphs = [
        DocumentParagraph(
            index=0,
            paragraph_type=ParagraphType.BODY,
            runs=runs,
        ),
    ]
    return DocumentStructure(
        filename="test.docx",
        paragraphs=paragraphs,
        total_paragraphs=1,
        total_words=4,
    )


@pytest.fixture
def tmp_dir():
    """Provide a temporary directory that is cleaned up after the test."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_docx_path(tmp_dir: Path) -> Path:
    """Create a small .docx with normal text, a heading, and a red placeholder."""
    doc = DocxDocument()

    doc.add_heading("Testdokument", level=1)

    p1 = doc.add_paragraph()
    run_normal = p1.add_run("Dies ist ein normaler Absatz mit ")
    run_normal.font.size = Pt(12)
    run_bold = p1.add_run("fettgedrucktem")
    run_bold.bold = True
    run_bold.font.size = Pt(12)
    run_end = p1.add_run(" Text.")
    run_end.font.size = Pt(12)

    p2 = doc.add_paragraph()
    run_placeholder = p2.add_run("[Platzhalter bitte ersetzen]")
    run_placeholder.font.color.rgb = RGBColor(255, 0, 0)

    p3 = doc.add_paragraph()
    run_italic = p3.add_run("Ein kursiver Satz.")
    run_italic.italic = True

    path = tmp_dir / "sample_formatted.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx_path(tmp_dir: Path) -> Path:
    """Create an empty .docx file."""
    doc = DocxDocument()
    path = tmp_dir / "empty.docx"
    doc.save(str(path))
    return path
